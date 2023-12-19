import os
import tkinter as tk
from tkinter import filedialog
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx.shared import Pt

class DocumentGenerator:
    def __init__(self, master):
        self.master = master
        self.master.title("Генератор Документов")
        self.master.geometry("500x650")

        self.headers_frame = tk.Frame(self.master, bg="#f0f0f0")
        self.headers_frame.pack(pady=10)

        self.format_frame = tk.Frame(self.master, bg="#f0f0f0")
        self.format_frame.pack(pady=10)

        self.output_frame = tk.Frame(self.master, bg="#f0f0f0")
        self.output_frame.pack(pady=10)

        self.font_frame = tk.Frame(self.master, bg="#f0f0f0")
        self.font_frame.pack(pady=10)

        # Путь к корневой папке
        self.root_path = os.path.dirname(__file__)

        # Путь к каталогу с шрифтами
        self.fonts_path = os.path.join(self.root_path, "fonts")

        # Получение списка шрифтов в каталоге
        self.available_fonts = [font_name[:-4] for font_name in os.listdir(self.fonts_path) if font_name.endswith('.ttf')]

        self.create_widgets()

    def create_widgets(self):
        # Заголовки
        tk.Label(self.headers_frame, text="Введите заголовки (каждый заголовок с новой строки):", font=("Helvetica", 12), bg="#f0f0f0").pack(pady=5)
        self.headers_text = tk.Text(self.headers_frame, height=5, width=30, font=("Helvetica", 12))
        self.headers_text.pack(padx=10)

        # Формат
        tk.Label(self.format_frame, text="Выберите формат:", font=("Helvetica", 12), bg="#f0f0f0").pack(pady=5)
        self.format_var = tk.StringVar()
        self.format_var.set(".docx")  # Формат по умолчанию
        formats = [".docx", ".pdf"]
        self.format_dropdown = tk.OptionMenu(self.format_frame, self.format_var, *formats)
        self.format_dropdown.config(font=("Helvetica", 12))
        self.format_dropdown.pack(padx=10)

        # Выходной файл
        tk.Label(self.output_frame, text="Имя генерируемого файла:", font=("Helvetica", 12), bg="#f0f0f0").pack(pady=5)
        self.output_var = tk.StringVar()
        self.output_entry = tk.Entry(self.output_frame, textvariable=self.output_var, font=("Helvetica", 12))
        self.output_entry.pack(padx=10)

        # Выбор места сохранения
        browse_button = tk.Button(self.output_frame, text="Обзор", command=self.browse_output_location, font=("Helvetica", 12), bg="#2196f3", fg="white")
        browse_button.pack(pady=5)

        # Выбор шрифта
        tk.Label(self.font_frame, text="Выберите шрифт:", font=("Helvetica", 12), bg="#f0f0f0").pack(pady=5)
        self.font_var = tk.StringVar()
        self.font_var.set(self.available_fonts[0])  # Шрифт по умолчанию
        self.font_dropdown = tk.OptionMenu(self.font_frame, self.font_var, *self.available_fonts)
        self.font_dropdown.config(font=("Helvetica", 12))
        self.font_dropdown.pack(pady=5)

        # Выбор размера шрифта
        tk.Label(self.font_frame, text="Выберите размер шрифта:", font=("Helvetica", 12), bg="#f0f0f0").pack(pady=5)
        self.font_size_var = tk.IntVar()
        self.font_size_var.set(12)  # Размер шрифта по умолчанию
        font_sizes = [10, 12, 14, 16, 18]
        self.font_size_dropdown = tk.OptionMenu(self.font_frame, self.font_size_var, *font_sizes)
        self.font_size_dropdown.config(font=("Helvetica", 12))
        self.font_size_dropdown.pack(pady=5)

        # Кнопки
        generate_button = tk.Button(self.master, text="Сгенерировать документ", command=self.generate_document, font=("Helvetica", 12), bg="#4caf50", fg="white")
        generate_button.pack(pady=10)

    def generate_document(self):
        headers = self.headers_text.get("1.0", tk.END).strip()
        format_selected = self.format_var.get()
        output_name = self.output_var.get()
        font_selected = self.font_var.get()
        font_size = self.font_size_var.get()

        if not headers or not output_name:
            tk.messagebox.showerror("Ошибка", "Заполните все поля.")
            return

        header_list = [header.strip() for header in headers.split('\n')]
        document = self.create_document(header_list, format_selected, output_name, font_selected, font_size)
        success = self.save_document(document, output_name, format_selected)

        if success:
            tk.messagebox.showinfo("Успех", f"Документ сохранен как {output_name + format_selected}")

    def create_document(self, headers, format_selected, output_name, font_selected, font_size):
        document = None

        if format_selected == ".docx":
            document = Document()

            for header in headers:
                # Создаем параграф для заголовка
                paragraph = document.add_paragraph()

                # Создаем объект заголовка уровня 1
                run = paragraph.add_run(header)

                # Устанавливаем шрифт и размер
                font = run.font
                font.name = font_selected
                font.size = Pt(font_size)

        elif format_selected == ".pdf":
            document = canvas.Canvas(output_name + format_selected, pagesize=letter)
            y_position = letter[1] - 50
            for header in headers:
                document.setFont(font_selected, font_size)
                document.drawString(100, y_position, header)
                y_position -= 20

        return document

    def save_document(self, document, output_name, format_selected):
        try:
            if format_selected == ".pdf":
                document.save()
            else:
                document.save(output_name + format_selected)
            return True
        except Exception as e:
            tk.messagebox.showerror("Ошибка", f"Ошибка сохранения документа: {str(e)}")
            return False

    def browse_output_location(self):
        output_location = filedialog.asksaveasfilename(defaultextension=".docx",
                                                       filetypes=[("Документы Word", "*.docx"),
                                                                  ("PDF-документы", "*.pdf")])
        self.output_var.set(output_location)

# Пример использования класса
def main():
    root = tk.Tk()
    app = DocumentGenerator(root)
    root.mainloop()

if __name__ == "__main__":
    main()